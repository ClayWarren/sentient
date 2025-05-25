"""
Advanced Features for Sentient - SOTA capabilities for 2025
Includes: File processing, Voice, Video, Memory, and Reasoning
"""

import os
import json
import logging
from typing import Dict, Any, List, Optional, Union
from dataclasses import dataclass
import time
from datetime import datetime
import numpy as np
import io
import base64

# Core dependencies
import torch
import pandas as pd
from PIL import Image
import cv2

# Document processing
import PyPDF2
from docx import Document
import openpyxl

# Voice processing
import speech_recognition as sr
import pyttsx3

# Vector storage and retrieval
import chromadb
from chromadb.config import Settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings

# Advanced reasoning
from typing import Tuple

logger = logging.getLogger(__name__)

@dataclass
class FileProcessingResult:
    """Result from file processing"""
    filename: str
    content: str
    metadata: Dict[str, Any]
    embeddings: Optional[List[float]] = None
    processing_time: float = 0.0

@dataclass
class VoiceResult:
    """Result from voice processing"""
    text: str
    confidence: float
    duration: float
    language: str = "en-US"

@dataclass
class VideoFrame:
    """Processed video frame data"""
    frame_number: int
    timestamp: float
    description: str
    objects_detected: List[str]
    frame_embedding: Optional[List[float]] = None

class AdvancedMemorySystem:
    """Persistent memory with vector search capabilities"""
    
    def __init__(self, persist_directory: str = "./data/memory"):
        self.persist_directory = persist_directory
        os.makedirs(persist_directory, exist_ok=True)
        
        # Initialize ChromaDB for vector storage
        self.chroma_client = chromadb.PersistentClient(
            path=persist_directory,
            settings=Settings(anonymized_telemetry=False)
        )
        
        # Create collections
        self.short_term = self.chroma_client.get_or_create_collection("short_term_memory")
        self.long_term = self.chroma_client.get_or_create_collection("long_term_memory")
        self.episodic = self.chroma_client.get_or_create_collection("episodic_memory")
        
        # Initialize embeddings model
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        logger.info("Advanced Memory System initialized with vector storage")
    
    def store_memory(self, content: str, memory_type: str = "short_term", 
                    metadata: Optional[Dict[str, Any]] = None) -> str:
        """Store memory with embeddings"""
        memory_id = f"{memory_type}_{int(time.time() * 1000)}"
        
        # Generate embedding
        embedding = self.embeddings.embed_query(content)
        
        # Select collection
        collection = getattr(self, memory_type, self.short_term)
        
        # Store in ChromaDB
        collection.add(
            embeddings=[embedding],
            documents=[content],
            metadatas=[metadata or {}],
            ids=[memory_id]
        )
        
        return memory_id
    
    def retrieve_memories(self, query: str, memory_type: str = "all", 
                         top_k: int = 5) -> List[Dict[str, Any]]:
        """Retrieve relevant memories using vector similarity"""
        query_embedding = self.embeddings.embed_query(query)
        
        results = []
        
        if memory_type == "all":
            collections = [self.short_term, self.long_term, self.episodic]
        else:
            collections = [getattr(self, memory_type, self.short_term)]
        
        for collection in collections:
            try:
                result = collection.query(
                    query_embeddings=[query_embedding],
                    n_results=min(top_k, collection.count())
                )
                
                for i in range(len(result['ids'][0])):
                    results.append({
                        'memory_id': result['ids'][0][i],
                        'content': result['documents'][0][i],
                        'metadata': result['metadatas'][0][i],
                        'distance': result['distances'][0][i],
                        'collection': collection.name
                    })
            except:
                continue
        
        # Sort by relevance (lower distance is better)
        results.sort(key=lambda x: x['distance'])
        
        return results[:top_k]
    
    def consolidate_memories(self):
        """Move important short-term memories to long-term storage"""
        # Get all short-term memories
        all_memories = self.short_term.get()
        
        if not all_memories['ids']:
            return
        
        # Simple importance scoring based on access frequency
        for i, memory_id in enumerate(all_memories['ids']):
            metadata = all_memories['metadatas'][i]
            access_count = metadata.get('access_count', 0)
            
            # Move to long-term if accessed frequently
            if access_count > 3:
                self.long_term.add(
                    embeddings=[all_memories['embeddings'][i]],
                    documents=[all_memories['documents'][i]],
                    metadatas=[metadata],
                    ids=[f"lt_{memory_id}"]
                )
                
                # Remove from short-term
                self.short_term.delete(ids=[memory_id])

class FileProcessor:
    """Advanced file processing with multiple format support"""
    
    def __init__(self, memory_system: Optional[AdvancedMemorySystem] = None):
        self.memory_system = memory_system
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
    
    def process_file(self, file_path: str, file_content: Optional[bytes] = None) -> FileProcessingResult:
        """Process various file types and extract content"""
        start_time = time.time()
        
        # Determine file type
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_content is None and os.path.exists(file_path):
                with open(file_path, 'rb') as f:
                    file_content = f.read()
            
            # Process based on file type
            if file_extension == '.pdf':
                content, metadata = self._process_pdf(file_content)
            elif file_extension in ['.doc', '.docx']:
                content, metadata = self._process_word(file_content)
            elif file_extension in ['.xls', '.xlsx']:
                content, metadata = self._process_excel(file_content)
            elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                content, metadata = self._process_image(file_content)
            elif file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css']:
                content = file_content.decode('utf-8', errors='ignore')
                metadata = {'type': 'text', 'encoding': 'utf-8'}
            elif file_extension == '.json':
                content, metadata = self._process_json(file_content)
            else:
                content = f"Unsupported file type: {file_extension}"
                metadata = {'type': 'unsupported'}
            
            # Store in memory if available
            embeddings = None
            if self.memory_system and len(content) > 0:
                memory_id = self.memory_system.store_memory(
                    content,
                    memory_type="long_term",
                    metadata={'filename': os.path.basename(file_path), **metadata}
                )
                metadata['memory_id'] = memory_id
            
            processing_time = time.time() - start_time
            
            return FileProcessingResult(
                filename=os.path.basename(file_path),
                content=content,
                metadata=metadata,
                embeddings=embeddings,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return FileProcessingResult(
                filename=os.path.basename(file_path),
                content=f"Error processing file: {str(e)}",
                metadata={'error': str(e)},
                processing_time=time.time() - start_time
            )
    
    def _process_pdf(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        metadata = {
            'type': 'pdf',
            'pages': len(pdf_reader.pages),
            'has_images': any(page.images for page in pdf_reader.pages)
        }
        
        return text.strip(), metadata
    
    def _process_word(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Word documents"""
        doc = Document(io.BytesIO(content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = {
            'type': 'word',
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables)
        }
        
        return text.strip(), metadata
    
    def _process_excel(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract data from Excel files"""
        df_dict = pd.read_excel(io.BytesIO(content), sheet_name=None)
        
        text_parts = []
        total_rows = 0
        
        for sheet_name, df in df_dict.items():
            text_parts.append(f"Sheet: {sheet_name}")
            text_parts.append(df.to_string())
            total_rows += len(df)
        
        metadata = {
            'type': 'excel',
            'sheets': len(df_dict),
            'total_rows': total_rows
        }
        
        return "\n\n".join(text_parts), metadata
    
    def _process_image(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process image and extract description"""
        image = Image.open(io.BytesIO(content))
        
        # Basic image analysis
        metadata = {
            'type': 'image',
            'format': image.format,
            'size': image.size,
            'mode': image.mode
        }
        
        # Convert to base64 for storage
        buffered = io.BytesIO()
        image.save(buffered, format=image.format or 'PNG')
        img_base64 = base64.b64encode(buffered.getvalue()).decode()
        
        description = f"Image: {metadata['format']} format, {metadata['size'][0]}x{metadata['size'][1]} pixels"
        
        return description, {**metadata, 'base64': img_base64}
    
    def _process_json(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process JSON files"""
        data = json.loads(content.decode('utf-8'))
        
        metadata = {
            'type': 'json',
            'keys': list(data.keys()) if isinstance(data, dict) else None,
            'length': len(data)
        }
        
        return json.dumps(data, indent=2), metadata

class VoiceProcessor:
    """Voice input/output processing"""
    
    def __init__(self):
        # Initialize speech recognition
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        
        # Initialize text-to-speech
        self.tts_engine = pyttsx3.init()
        self._setup_tts()
        
        logger.info("Voice Processor initialized")
    
    def _setup_tts(self):
        """Configure TTS engine"""
        voices = self.tts_engine.getProperty('voices')
        # Try to use a female voice if available
        for voice in voices:
            if 'female' in voice.name.lower():
                self.tts_engine.setProperty('voice', voice.id)
                break
        
        self.tts_engine.setProperty('rate', 180)  # Speed
        self.tts_engine.setProperty('volume', 0.9)  # Volume
    
    def listen(self, timeout: float = 5.0, phrase_limit: float = 10.0) -> VoiceResult:
        """Listen for voice input"""
        start_time = time.time()
        
        try:
            with self.microphone as source:
                # Adjust for ambient noise
                self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
                
                # Listen for audio
                audio = self.recognizer.listen(
                    source,
                    timeout=timeout,
                    phrase_time_limit=phrase_limit
                )
                
                # Recognize speech
                text = self.recognizer.recognize_google(audio)
                confidence = 0.9  # Google doesn't provide confidence scores
                
                duration = time.time() - start_time
                
                return VoiceResult(
                    text=text,
                    confidence=confidence,
                    duration=duration
                )
                
        except sr.WaitTimeoutError:
            return VoiceResult(
                text="",
                confidence=0.0,
                duration=time.time() - start_time
            )
        except sr.UnknownValueError:
            return VoiceResult(
                text="[Could not understand audio]",
                confidence=0.0,
                duration=time.time() - start_time
            )
        except Exception as e:
            logger.error(f"Voice recognition error: {e}")
            return VoiceResult(
                text=f"[Error: {str(e)}]",
                confidence=0.0,
                duration=time.time() - start_time
            )
    
    def speak(self, text: str, wait: bool = True):
        """Convert text to speech"""
        try:
            self.tts_engine.say(text)
            if wait:
                self.tts_engine.runAndWait()
            else:
                # Run in background
                import threading
                thread = threading.Thread(target=self.tts_engine.runAndWait)
                thread.daemon = True
                thread.start()
        except Exception as e:
            logger.error(f"TTS error: {e}")

class VideoProcessor:
    """Real-time video processing and analysis"""
    
    def __init__(self):
        self.cap = None
        self.is_processing = False
        
        # Load pre-trained models for object detection
        self.face_cascade = cv2.CascadeClassifier(
            cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
        )
        
        logger.info("Video Processor initialized")
    
    def start_capture(self, source: Union[int, str] = 0):
        """Start video capture from camera or file"""
        self.cap = cv2.VideoCapture(source)
        self.is_processing = True
        
        if not self.cap.isOpened():
            raise RuntimeError("Failed to open video source")
    
    def process_frame(self, frame: np.ndarray) -> VideoFrame:
        """Process a single video frame"""
        frame_gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        
        # Detect faces
        faces = self.face_cascade.detectMultiScale(frame_gray, 1.1, 4)
        
        objects_detected = []
        if len(faces) > 0:
            objects_detected.append(f"{len(faces)} face(s)")
        
        # Basic frame description
        height, width = frame.shape[:2]
        description = f"Frame {width}x{height}, detected: {', '.join(objects_detected) if objects_detected else 'nothing'}"
        
        return VideoFrame(
            frame_number=0,
            timestamp=time.time(),
            description=description,
            objects_detected=objects_detected
        )
    
    def get_frame(self) -> Optional[Tuple[np.ndarray, VideoFrame]]:
        """Get next frame from video source"""
        if self.cap is None or not self.cap.isOpened():
            return None
        
        ret, frame = self.cap.read()
        if not ret:
            return None
        
        frame_data = self.process_frame(frame)
        return frame, frame_data
    
    def stop_capture(self):
        """Stop video capture"""
        self.is_processing = False
        if self.cap:
            self.cap.release()
            self.cap = None

class ChainOfThoughtReasoner:
    """Advanced reasoning with explicit thought chains"""
    
    def __init__(self, base_model):
        self.base_model = base_model
        self.thought_steps = []
    
    def reason(self, problem: str, max_steps: int = 5) -> Dict[str, Any]:
        """Perform step-by-step reasoning"""
        self.thought_steps = []
        
        current_context = problem
        final_answer = None
        
        for step in range(max_steps):
            # Generate thought step
            thought_prompt = f"""Step {step + 1} of solving this problem:
Problem: {problem}
Previous steps: {self._format_previous_steps()}
Current context: {current_context}

What is the next logical step in solving this? Think step by step."""
            
            # Generate thought
            thought_result = self.base_model.generate(
                thought_prompt,
                max_tokens=150
            )
            
            thought = thought_result.text
            self.thought_steps.append({
                'step': step + 1,
                'thought': thought,
                'confidence': thought_result.confidence
            })
            
            # Check if we have an answer
            if any(indicator in thought.lower() for indicator in ['therefore', 'the answer is', 'in conclusion']):
                final_answer = thought
                break
            
            current_context = thought
        
        # Generate final answer if not found
        if final_answer is None:
            answer_prompt = f"""Based on the following reasoning steps, what is the final answer?
Problem: {problem}
Reasoning steps:
{self._format_all_steps()}

Final answer:"""
            
            answer_result = self.base_model.generate(answer_prompt, max_tokens=100)
            final_answer = answer_result.text
        
        return {
            'problem': problem,
            'thought_chain': self.thought_steps,
            'final_answer': final_answer,
            'total_steps': len(self.thought_steps)
        }
    
    def _format_previous_steps(self) -> str:
        """Format previous thought steps"""
        if not self.thought_steps:
            return "None"
        
        return "\n".join([f"Step {s['step']}: {s['thought']}" for s in self.thought_steps[-3:]])
    
    def _format_all_steps(self) -> str:
        """Format all thought steps"""
        return "\n".join([f"Step {s['step']}: {s['thought']}" for s in self.thought_steps])

# Initialize advanced features
def initialize_advanced_features(consciousness_ai=None) -> Dict[str, Any]:
    """Initialize all advanced feature modules"""
    
    # Initialize memory system
    memory = AdvancedMemorySystem()
    
    # Initialize processors
    file_processor = FileProcessor(memory_system=memory)
    voice_processor = VoiceProcessor()
    video_processor = VideoProcessor()
    
    # Initialize chain-of-thought reasoner if AI is provided
    reasoner = ChainOfThoughtReasoner(consciousness_ai) if consciousness_ai else None
    
    logger.info("✨ Advanced features initialized successfully")
    
    return {
        'memory': memory,
        'file_processor': file_processor,
        'voice_processor': voice_processor,
        'video_processor': video_processor,
        'reasoner': reasoner
    }